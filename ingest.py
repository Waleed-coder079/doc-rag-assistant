import os
import re
import json
import uuid
from pathlib import Path
from pdfminer.high_level import extract_text as pdf_extract_text
from llama_index.core import Document
from docx import Document as DocxReader
import tabula
import pandas as pd
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl


# ========== Table-specific helpers ==========

# Regex patterns for heading detection. These are used to spot likely section
# titles that often precede tables in the documents.
# - HEADING_RE_NUMBER: matches numbered headings like "3.2 VAT Rates"
HEADING_RE_NUMBER = re.compile(r'^\s*\d+(?:\.\d+)*\s+\S+')
# - HEADING_RE_UPPER: matches ALL-CAPS headings commonly used as section headers
HEADING_RE_UPPER = re.compile(r'^[A-Z0-9][A-Z0-9 \-/&()]{3,}$')
# - HEADING_RE_TITLE_CASE: matches Title Case headings like "Electronic Filing Methods"
HEADING_RE_TITLE_CASE = re.compile(r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+$')

def clean_table_cell(val: str) -> str:
    """Clean individual table cell without destroying column boundaries."""
    if val is None:
        return ""
    # Remove line breaks inside cell, trim, collapse internal multi-spaces
    v = str(val).replace('\r', ' ').replace('\n', ' ')
    v = re.sub(r'[ \t]+', ' ', v)
    return v.strip()

def is_heading_candidate(line: str) -> bool:
    """Detect if a line looks like a heading/title.
    
    Based on document structure:
    - First letter is always capital
    - Often Title Case (first letter of each word capitalized)
    - Headings are surrounded by newlines (handled by caller)
    - No colon endings in your documents
    """
    l = line.strip()
    if not l:
        return False
    
    # Length constraints: headings are typically 3-150 chars
    if len(l) < 3 or len(l) > 150:
        return False
    
    # Must start with capital letter
    if not l[0].isupper():
        return False
    
    # Pattern 1: Numbered sections (e.g., "3.2 VAT Return Submission")
    if HEADING_RE_NUMBER.match(l):
        return True
    
    # Pattern 2: ALL CAPS (e.g., "VAT FILING REQUIREMENTS")
    if HEADING_RE_UPPER.match(l):
        return True
    
    # Pattern 3: Title Case - each word starts with capital
    # (e.g., "Electronic Filing Methods", "Country Filing Obligations")
    words = l.split()
    if 2 <= len(words) <= 10:  # Reasonable heading length
        # Check if most words start with capital (Title Case pattern)
        capitalized_words = sum(1 for w in words if w[0].isupper())
        if capitalized_words >= len(words) * 0.75:  # At least 75% words capitalized
            return True
    
    # Pattern 4: Short capitalized phrase (2-8 words, starts with capital)
    if len(words) <= 8:
        capital_ratio = sum(1 for c in l if c.isupper()) / max(1, len(l))
        if capital_ratio > 0.15:  # At least 15% capitals
            return True
    
    return False

def score_heading(line: str) -> int:
    """Score heading strength for prioritization.
    
    Higher scores indicate stronger heading signals.
    """
    if not line:
        return 0
    
    l = line.strip()
    score = 0
    
    # Numbered sections are strongest (e.g., "3.2 VAT Filing")
    if HEADING_RE_NUMBER.match(l):
        score += 5
    
    # ALL CAPS headings are very strong
    if HEADING_RE_UPPER.match(l):
        score += 4
    
    # Title Case pattern (e.g., "Electronic Filing Methods")
    if HEADING_RE_TITLE_CASE.match(l):
        score += 3
    
    # Bonus for concise headings (better than long ones)
    word_count = len(l.split())
    if 2 <= word_count <= 6:
        score += 2
    elif word_count <= 8:
        score += 1
    
    # Bonus for high capital letter ratio (indicates heading style)
    capital_ratio = sum(1 for c in l if c.isupper()) / max(1, len(l))
    if capital_ratio > 0.3:
        score += 1
    
    return score

def find_table_title_from_lines(lines: list[str], max_lookback: int = 15) -> tuple[str | None, int]:
    """Find the best heading candidate from a list of preceding lines.
    
    Headings in your documents:
    - Are surrounded by blank lines (newlines before/after)
    - Start with capital letter
    - Often Title Case (each word capitalized)
    
    Returns (title, confidence_score).
    """
    best_title = None
    best_score = 0
    
    # Filter out empty lines first
    non_empty_lines = [l for l in lines if l.strip()]
    
    # Look at last N non-empty lines (reversed, most recent first)
    lookback_lines = list(reversed(non_empty_lines[-max_lookback:]))
    
    for idx, line in enumerate(lookback_lines):
        line_stripped = line.strip()
        
        if is_heading_candidate(line_stripped):
            s = score_heading(line_stripped)
            
            # Bonus: Check if line is isolated (has empty lines around it)
            # This indicates it's likely a heading, not body text
            is_isolated = False
            
            # For the original lines (with empties), check context
            if line_stripped in lines:
                orig_idx = len(lines) - 1 - lines[::-1].index(line_stripped)
                
                # Check if preceded by empty line
                has_blank_before = (orig_idx > 0 and not lines[orig_idx - 1].strip())
                
                # Check if followed by empty line
                has_blank_after = (orig_idx < len(lines) - 1 and not lines[orig_idx + 1].strip())
                
                if has_blank_before or has_blank_after:
                    is_isolated = True
                    s += 2  # Bonus for isolated heading
            
            if s > best_score:
                best_score = s
                best_title = line_stripped
            
            # Early exit if very strong heading (score >= 6)
            if best_score >= 6:
                break
    
    return best_title, best_score

def serialize_table_dataframe(df: pd.DataFrame, source_path: str, title: str | None = None, page: int | None = None) -> tuple[str, dict] | None:
    """Convert DataFrame to searchable text + structured JSON metadata."""
    if df is None or df.empty:
        return None
    
    df = df.fillna("")
    
    # Normalize column headers
    cols = []
    for i, c in enumerate(df.columns):
        c_clean = clean_table_cell(str(c))
        cols.append(c_clean if c_clean else f"col_{i+1}")
    df.columns = cols
    
    # Clean each cell individually to preserve column boundaries
    for col in df.columns:
        df[col] = df[col].apply(clean_table_cell)
    
    # Convert to structured JSON
    rows = df.to_dict(orient="records")
    
    # Create searchable text representation (pipe-delimited)
    header_line = " | ".join(cols)
    row_lines = [" | ".join(str(r.get(c, "")) for c in cols) for r in rows]
    
    # Build text block with title
    if title:
        text_block = f"TABLE TITLE: {title}\n\n"
    else:
        text_block = ""
    
    text_block += f"TABLE from {os.path.basename(source_path)}\n"
    text_block += f"{header_line}\n"
    text_block += "\n".join(row_lines)
    
    # Metadata with full JSON structure
    metadata = {
        "type": "table",
        "table_title": title,
        "table_data": {
            "columns": cols,
            "rows": rows,
            "row_count": len(rows),
            "col_count": len(cols)
        },
        "page": page,
        "source_file": os.path.basename(source_path)
    }
    
    return text_block, metadata

# ========== Text cleaning functions ==========

def clean_text(text: str) -> str:
    """Normalize and clean text:
       - normalize CRLF -> LF and form-feeds
       - remove long dot leaders (3+ dots) and long dashes/equals
       - collapse repeated blank/whitespace-only lines into a single newline
       - collapse runs of spaces/tabs but preserve single newlines
       - strip each line's ends
    """
    if not text:
        return ""

    text = text.replace('\r\n', '\n').replace('\r', '\n')
    text = text.replace('\x0c', '\n')

    text = re.sub(r'\.{3,}', ' ', text)
    text = re.sub(r'[-=]{3,}', ' ', text)
    text = re.sub(r'(?:\n\s*){2,}', '\n', text)
    text = re.sub(r'[ \t]+', ' ', text)

    lines = [ln.strip() for ln in text.split('\n')]
    text = '\n'.join(lines).strip()

    return text


def _dedupe_pdf_header(text: str, lookahead_lines: int = 10) -> str:
    """Remove duplicate headers in PDFs."""
    if not text:
        return text

    lines = text.split('\n')
    first_idx = None
    for i, l in enumerate(lines):
        if l.strip():
            first_idx = i
            break
    if first_idx is None:
        return text

    title_line = lines[first_idx].strip()
    header_candidate = None

    for j in range(first_idx + 1, min(len(lines), first_idx + 1 + lookahead_lines)):
        cand = lines[j].strip()
        if ':' in cand and title_line and title_line in cand:
            header_candidate = cand
            break

    seen_title = False
    seen_header = False
    out_lines = []
    for l in lines:
        s = l.strip()
        if not s:
            out_lines.append('')
            continue

        if s == title_line:
            if not seen_title:
                out_lines.append(s)
                seen_title = True
            continue
        elif header_candidate and s == header_candidate:
            if not seen_header:
                out_lines.append(s)
                seen_header = True
            continue
        else:
            out_lines.append(l)

    return '\n'.join(out_lines)


def _remove_pdf_page_numbers(text: str) -> str:
    """
    Remove isolated page numbers from PDF text.
    Matches lines that contain only digits (with optional spaces).
    Example: '1\\n   2\\n  3\\n' -> '\\n\\n\\n'
    """
    if not text:
        return text

    text = re.sub(r'^\s*\d+\s*$', '', text, flags=re.MULTILINE)
    text = re.sub(r'(?:\n\s*){2,}', '\n', text)
    return text


def _extract_table_of_contents(text: str, min_lines: int = 2):
    """
    Extract the 'Table of Contents' block from text if present.
    Returns (toc_text, cleaned_text). If no TOC found, returns (None, original_text).

    Logic:
      - Find a heading line like 'Table of Contents' or 'Contents'
      - From the next line onward, collect contiguous 'TOC-like' lines:
        * lines that end with a page number (e.g. "Overview 3")
        * short numbered lines (e.g. "1." or "2.1. Section")
        * short lines with dot-leaders or that start with a digit
      - Stop when we encounter a couple of consecutive non-TOC lines.
      - Only treat as TOC if we captured at least `min_lines` non-empty lines.
    """
    if not text:
        return None, text

    lines = text.split('\n')
    heading_idx = None
    for i, l in enumerate(lines):
        # Matches lines like 'Table of Contents' (case-insensitive)
        if re.match(r'^\s*table of (contents|content)s?\b', l, flags=re.I):
            heading_idx = i
            break
    if heading_idx is None:
        # Fallback: heading might be simply 'Contents'
        for i, l in enumerate(lines):
            if re.match(r'^\s*contents\b', l, flags=re.I):
                heading_idx = i
                break
    if heading_idx is None:
        return None, text

    toc_lines = []
    non_toc_count = 0
    last_idx = heading_idx

    for i in range(heading_idx + 1, len(lines)):
        raw = lines[i]
        s = raw.strip()

        if s == "":
            # accept blank line as part of TOC (reset non_toc_count)
            toc_lines.append('')
            last_idx = i
            non_toc_count = 0
            continue

        # Patterns that indicate a TOC entry. We compute several heuristics:
        # - ends_with_page: entries like 'Overview 3' (trailing page number)
        # - only_number_dot: lines like '1.' or '2.1.' which look like TOC indices
        # - starts_with_number: '1. Introduction' style
        # - has_dot_leaders: visual dot leaders '.....' before page numbers
        ends_with_page = re.search(r'\b\d{1,3}\s*$', s) is not None
        only_number_dot = re.match(r'^\d+(\.\d+)*\.?$', s) is not None
        starts_with_number = re.match(r'^\d+(\.\d+)*\s+', s) is not None
        has_dot_leaders = re.search(r'\.{2,}|·|\t', raw) is not None
        short_line = len(s) <= 120

        if ends_with_page or only_number_dot or (short_line and (starts_with_number or has_dot_leaders)):
            toc_lines.append(s)
            last_idx = i
            non_toc_count = 0
            continue
        else:
            # allow a limited number of "non-TOC" lines to be tolerant of noisy PDFs
            non_toc_count += 1
            if non_toc_count >= 2:
                break
            # allow a single short non-matching line to be included (for noisy PDFs)
            if len(s) < 40:
                toc_lines.append(s)
                last_idx = i
                continue
            else:
                continue

    # drop trailing blank lines from toc_lines
    while toc_lines and toc_lines[-1].strip() == "":
        toc_lines.pop()

    # require at least min_lines TOC lines (non-empty) to avoid false positives
    # Require a minimum number of TOC-like lines to avoid false positives
    if len([ln for ln in toc_lines if ln.strip() != ""]) < min_lines:
        return None, text

    toc_block = '\n'.join([lines[heading_idx].strip()] + toc_lines).strip()

    # remove the TOC block lines from the main text
    cleaned_lines = lines[:heading_idx] + lines[last_idx + 1 :]
    cleaned_text = '\n'.join(cleaned_lines).strip()

    return toc_block, cleaned_text


def extract_text_with_pages(filepath: str):
    """Extract text with page-like splits (PDF vs DOCX)."""
    ext = Path(filepath).suffix.lower()

    if ext == ".pdf":
        text = pdf_extract_text(filepath)
        pages = text.split("\f") if text else []
        full_text = "\n".join(pages)

        # Clean and normalize
        full_text = clean_text(full_text)

        # Remove duplicate headers
        full_text = _dedupe_pdf_header(full_text)

        # Remove page numbers
        full_text = _remove_pdf_page_numbers(full_text)

        # Extract TOC (if present) and remove it from body
        toc_text, full_text = _extract_table_of_contents(full_text)

        results = []
        if full_text.strip():
            results.append((full_text, {}))

        if toc_text:
            results.append((toc_text, {"type": "table_of_contents"}))

        # Extract tables page-by-page to capture context for title detection
        try:
            for page_num in range(1, len(pages) + 1):
                try:
                    page_tables = tabula.read_pdf(filepath, pages=str(page_num), multiple_tables=True)
                except Exception:
                    continue
                
                if not page_tables:
                    continue
                
                # Get cleaned text lines from this page for title detection
                page_text = pages[page_num - 1]
                page_cleaned = clean_text(page_text)
                page_lines = [ln for ln in page_cleaned.split('\n') if ln.strip()]
                
                # Process each table found on this page
                for table_df in page_tables:
                    if isinstance(table_df, pd.DataFrame) and not table_df.empty:
                        # Find title from preceding lines on the same page
                        title, confidence = find_table_title_from_lines(page_lines)
                        
                        # Serialize with title and page info
                        serialized = serialize_table_dataframe(
                            table_df, 
                            filepath, 
                            title=title, 
                            page=page_num
                        )
                        
                        if serialized:
                            table_text, table_meta = serialized
                            table_meta['title_confidence'] = confidence
                            results.append((table_text, table_meta))
        except Exception as e:
            print(f"Table extraction failed for PDF {filepath}: {e}")

        return results

    elif ext == ".docx":
        doc = DocxReader(filepath)
        
        # Build ordered sequence of elements (paragraphs + tables) to preserve document structure
        elements = []
        paragraph_iter = iter(doc.paragraphs)
        table_iter = iter(doc.tables)
        
        # Walk document body in order
        for child in doc.element.body.iterchildren():
            if isinstance(child, CT_P):
                para = next(paragraph_iter)
                elements.append(("paragraph", para))
            elif isinstance(child, CT_Tbl):
                tbl = next(table_iter)
                elements.append(("table", tbl))
        
        results = []
        body_paragraphs = []
        
        # Process elements in order
        for idx, (elem_type, elem_obj) in enumerate(elements):
            if elem_type == "paragraph":
                para_text = elem_obj.text.strip()
                if para_text:
                    body_paragraphs.append(para_text)
            
            elif elem_type == "table":
                # Look backward to find title from preceding paragraphs
                preceding_lines = []
                for prev_idx in range(idx - 1, max(-1, idx - 20), -1):
                    if elements[prev_idx][0] == "paragraph":
                        prev_text = elements[prev_idx][1].text.strip()
                        if prev_text:
                            preceding_lines.insert(0, prev_text)
                    else:
                        # Stop at previous table
                        break
                
                # Find best title candidate
                title, confidence = find_table_title_from_lines(preceding_lines)
                
                # Extract table data
                rows = []
                for row in elem_obj.rows:
                    cells = [clean_table_cell(cell.text) for cell in row.cells]
                    rows.append(cells)
                
                if rows:
                    try:
                        # First row as header if multiple rows
                        if len(rows) > 1:
                            df = pd.DataFrame(rows[1:], columns=rows[0])
                        else:
                            df = pd.DataFrame(rows)
                    except Exception:
                        df = pd.DataFrame(rows)
                    
                    # Serialize with title
                    serialized = serialize_table_dataframe(df, filepath, title=title)
                    if serialized:
                        table_text, table_meta = serialized
                        table_meta['title_confidence'] = confidence
                        results.append((table_text, table_meta))
        
        # Add main body text (all paragraphs combined)
        if body_paragraphs:
            body_text = clean_text("\n".join(body_paragraphs))
            if body_text.strip():
                results.insert(0, (body_text, {}))
        
        return results

    else:
        raise ValueError(f"Unsupported file type: {ext}")


def process_document(filepath: str):
    """Process PDF or DOCX into LlamaIndex Documents with metadata."""
    records = []
    chunks = extract_text_with_pages(filepath)
    for chunk_text, extra_meta in chunks:
        if not chunk_text or not chunk_text.strip():
            continue
        meta = {
            "doc_id": str(uuid.uuid4()),
            "file_name": os.path.basename(filepath),
            "title": Path(filepath).stem,
        }
        meta.update(extra_meta)
        doc = Document(text=chunk_text, metadata=meta)
        records.append(doc)
    return records



def ingest(data_dir: str, out_file: str):
    Path(out_file).parent.mkdir(parents=True, exist_ok=True)
    all_records = []

    for file in Path(data_dir).glob("*"):
        # Only process PDFs and DOCX files for this project.
        if file.suffix.lower() in [".pdf", ".docx"]:
            recs = process_document(str(file))
        else:
            print(f"Skipping unsupported file type: {file}")
            continue

        print(f"Ingested {len(recs)} records from {file}")
        all_records.extend(recs)

    with open(out_file, "w", encoding="utf-8") as f:
        for rec in all_records:
            if isinstance(rec, Document):
                f.write(json.dumps({"text": rec.text, "metadata": rec.metadata}, ensure_ascii=False) + "\n")
            else:
                f.write(json.dumps(rec, ensure_ascii=False) + "\n")

    print(f"✅ Saved {len(all_records)} records to {out_file}")


if __name__ == "__main__":
    # Default run when executed directly
    INPUT_DIR = "input"
    OUTPUT_FILE = "ing_out_split_in/docs.jsonl"
    ingest(INPUT_DIR, OUTPUT_FILE)

